"""
Step 7: Generate the monthly business report as a .docx, mirroring the layout,
wording and logic of the reference report (Report_05-2026.docx).

Numbers/tables are computed deterministically; the analytical prose is written
by the LLM (see src/llm/report_writer.py).
"""
import re
from pathlib import Path
from typing import Dict, Any

from .base import BaseStep, StepResult
from .. import config


def _shade(cell, hex_color="E1E6F0"):
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


class Step7GenerateReport(BaseStep):
    @property
    def name(self) -> str:
        return "Step 7: Generate Report (.docx)"

    @property
    def description(self) -> str:
        return "Write the analytical monthly business report as a Word document"

    def execute(self, month: int, year: int, context: Dict[str, Any]) -> Dict[str, Any]:
        self.log(f"Generating .docx report for {config.get_month_name(month)} {year}")
        config.REPORT_DIR.mkdir(parents=True, exist_ok=True)

        from ..llm.report_writer import build_report_data, generate_narrative
        from ..llm import llm_available

        data = build_report_data(context, month, year, self.data_source)
        if llm_available():
            self.log("Writing analytical narrative via LLM...")
        else:
            self.log("LLM not configured — rendering tables with minimal text", "warn")
        narrative = generate_narrative(data)

        out = config.REPORT_DIR / f"Report thang {month} nam {year}.docx"
        try:
            self._render(out, data, narrative, month, year)
        except PermissionError:
            for n in range(1, 100):
                alt = config.REPORT_DIR / f"Report thang {month} nam {year} ({n}).docx"
                if not alt.exists():
                    break
            self.log(f"{out.name} đang mở/khoá; ghi {alt.name}", "warn")
            self._render(alt, data, narrative, month, year)
            out = alt

        context["report_path"] = str(out)
        context["pdf_path"] = str(out)  # backward-compat for callers
        self.log(f"Report saved to {out}", "success")
        return StepResult(success=True, data={"report_path": str(out)},
                          message=f"Report generated: {out.name}").__dict__

    # ------------------------------------------------------------------ #
    def _render(self, path: Path, data: dict, nar: dict, month: int, year: int):
        from docx import Document
        from docx.shared import Pt, RGBColor
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        meta = data["meta"]
        yy = year % 100
        prev3, cur3, next3 = meta["prev_month"][:3], meta["report_month"][:3], meta["next_month"][:3]

        BLUE = RGBColor(0x1F, 0x4E, 0x79)   # primary accent (titles, headings, table header)
        GREY = RGBColor(0x59, 0x59, 0x59)   # subtitle / meta text
        WHITE = RGBColor(0xFF, 0xFF, 0xFF)

        doc = Document()
        doc.styles["Normal"].font.name = "Calibri"
        doc.styles["Normal"].font.size = Pt(10)

        def para(text, bold=False, size=None, space_after=4, color=None):
            p = doc.add_paragraph()
            r = p.add_run(text)
            r.bold = bold
            if size:
                r.font.size = Pt(size)
            if color is not None:
                r.font.color.rgb = color
            p.paragraph_format.space_after = Pt(space_after)
            return p

        def subhead(text):
            return para(text, bold=True, size=11, color=BLUE, space_after=3)

        def heading(text):
            h = doc.add_heading(text, level=1)
            for r in h.runs:
                r.font.color.rgb = BLUE
                r.bold = True
            return h

        def label_para(text):
            """Bold the leading 'Label.' / 'Read:' part, rest normal."""
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            m = re.match(r"^([^.:]{1,40}[.:])\s*(.*)$", text, re.DOTALL)
            if m:
                p.add_run(m.group(1) + " ").bold = True
                p.add_run(m.group(2))
            else:
                p.add_run(text)
            return p

        def bullets(items):
            for it in (items or []):
                if str(it).strip():
                    doc.add_paragraph(str(it), style="List Bullet")

        def table(headers, rows, widths=None):
            t = doc.add_table(rows=1, cols=len(headers))
            t.style = "Table Grid"
            for j, h in enumerate(headers):
                c = t.rows[0].cells[j]
                c.text = ""
                run = c.paragraphs[0].add_run(h)
                run.bold = True
                run.font.size = Pt(8.5)
                run.font.color.rgb = WHITE
                _shade(c, "1F4E79")
            for row in rows:
                cells = t.add_row().cells
                for j, val in enumerate(row):
                    cells[j].text = ""
                    rr = cells[j].paragraphs[0].add_run("" if val is None else str(val))
                    rr.font.size = Pt(8.5)
            return t

        # --- Header ---
        para(meta["report_code"], bold=True, size=18, space_after=2, color=BLUE)
        para(f"Cash Loan (CL) monthly business report  •  Reporting month: "
             f"{meta['report_month']} {year}  •  Latest actual: {cur3}-{yy}  •  "
             f"Forecast month: {next3}-{yy}", size=9, space_after=2, color=GREY)
        para("Data basis. Figures are drawn from the Annual Plan / KPI sheet and the "
             "Initiative Tracker (status as of end-" + f"{meta['report_month']} {year}). "
             "Structure follows the fixed four-part format; each section is split into "
             "Structural ceiling unlocks and Incremental acquisition / retention improvements.",
             size=8.5, space_after=8, color=GREY)

        # --- KPI snapshot ---
        subhead("KPI snapshot — Cash Loan funnel")
        table(
            ["Metric (Cash Loan)", "Unit", f"{prev3}-{yy}", f"{cur3}-{yy}", "MoM", f"{cur3} Plan", "vs Plan"],
            [[r["metric"], r["unit"], r["prev"], r["cur"], r["mom"], r["plan"], r["vs_plan"]]
             for r in data["funnel"]],
        )
        if nar.get("kpi_read"):
            label_para("Read: " + nar["kpi_read"])
        else:
            para("", space_after=4)

        # Movement check (Rule 7) — deterministic real-vs-mechanical flags
        if data.get("movement"):
            label_para("Movement check: " + " ".join(data["movement"]))

        # --- 1. Performance Overview ---
        heading("1. Performance Overview and Analysis")
        for key in ("headline", "mom_bridge", "plan_miss_bridge"):
            if nar.get(key):
                label_para(nar[key])
        subhead("Structural ceiling unlocks")
        bullets(nar.get("structural_bullets") or self._fallback_bullets(data["structural"]))
        subhead("Incremental acquisition / retention improvements")
        bullets(nar.get("incremental_bullets") or self._fallback_bullets(data["incremental"]))

        # --- 2. Next month run-rate ---
        heading(f"2. Next Month ({meta['next_month']}) Run-Rate and Key Initiatives")
        if nar.get("runrate"):
            label_para(nar["runrate"])
        elif data["forecast"]["disb_next"] is not None:
            label_para(f"Run-rate. {meta['next_month']} Disbursement run-rate is "
                       f"~{data['forecast']['disb_next']:,} VNDm.")
        if nar.get("runrate_structural"):
            subhead("Structural ceiling unlocks")
            bullets(nar["runrate_structural"])
        if nar.get("runrate_incremental"):
            subhead("Incremental acquisition / retention improvements")
            bullets(nar["runrate_incremental"])

        # --- 3. Top priorities ---
        heading(f"3. Top Priorities for Next Month ({meta['next_month']})")
        tp = nar.get("top_priorities") or []
        if tp:
            table(["#", "Objective / lever", "Initiative & owner", "Target outcome & why it matters"],
                  [[i + 1, r.get("objective", ""), r.get("initiative_owner", ""), r.get("target_why", "")]
                   for i, r in enumerate(tp)])
        else:
            table(["#", "Initiative", "Status", "Timing", "Expected impact"],
                  [[i + 1, r["name"], r["status"], r["timing"], r["expected_impact"]]
                   for i, r in enumerate(data["top_priorities"])])

        # --- 4. Annual plan progress ---
        heading("4. Progress Toward Annual Plan")
        subhead("Structural ceiling unlocks")
        a_s = nar.get("annual_structural") or []
        if a_s:
            table(["Lever", "What landed / what's next", "Confidence", "Status"],
                  [[r.get("lever", ""), r.get("landed_next", ""), r.get("confidence", ""), r.get("status", "")]
                   for r in a_s])
        else:
            table(["No", "Initiative", "Status", "Details"],
                  [[r["no"], r["name"], r["status"], r["details"]] for r in data["structural"]])
        subhead("Incremental acquisition / retention improvements")
        a_i = nar.get("annual_incremental") or []
        if a_i:
            table(["Lever", "What landed / what's next", "Impact", "Status"],
                  [[r.get("lever", ""), r.get("landed_next", ""), r.get("impact", ""), r.get("status", "")]
                   for r in a_i])
        else:
            table(["No", "Initiative", "Status", "Details"],
                  [[r["no"], r["name"], r["status"], r["details"]] for r in data["incremental"]])

        subhead("YTD vs plan and FY outlook")
        if nar.get("ytd_outlook"):
            para(nar["ytd_outlook"])
        else:
            ytd = data["ytd"]
            if ytd["disb_ytd_actual"] is not None:
                para(f"YTD Disbursement is {ytd['disb_ytd_actual']:,} VNDm vs "
                     f"{ytd['disb_ytd_plan']:,} plan; FY plan {ytd['fy_plan']:,} VNDm.")
        # Escalations — deterministic (guaranteed), highest impact first.
        esc = data.get("escalations") or []
        if esc:
            subhead("Escalations — high-impact initiatives delayed / deprioritized")
            table(["Initiative", "Owner", "Status", "Impact", "New timing", "Note"],
                  [[e["name"], e["pic"], e["status"], e["expected_impact"],
                    (e["new_timing"] or "-"), e["details"]] for e in esc])

        subhead("Risks, dependencies & escalations")
        bullets(nar.get("risks"))

        # To-verify: source conflicts (Rule 3) — deterministic
        conflicts = data.get("conflicts") or []
        if conflicts:
            subhead("To verify — source conflicts (email vs actuals)")
            bullets([
                f"{c['metric']}: email"
                + (f" ({c['owner']})" if c.get("owner") else "")
                + f" says {c['email_value']} vs actuals sheet {c['sheet_value']} — reconcile."
                for c in conflicts
            ])

        doc.save(str(path))

    @staticmethod
    def _fallback_bullets(items):
        out = []
        for it in items:
            s = f"{it['name']} — status {it['status']}"
            if it["details"]:
                s += f"; {it['details']}"
            out.append(s)
        return out
